"""
Google Drive Ingestion Service
Syncs files from Google Drive to database with multi-user support
"""
from typing import Optional, List, Dict
from datetime import datetime, timezone
from sqlalchemy.orm import Session
from google_drive import GoogleDriveClient
from models.metadata import (
    Document, 
    ProcessingQueue, 
    SyncCheckpoint, 
    TaskType, 
    ProcessingStatus,
    Tag,
    DocumentTag,
    ContentType  # ADDED THIS IMPORT
)
from tagging import SimpleTagger
import hashlib
import config
import os

# Add conditional imports for text extraction libraries
try:
    import PyPDF2
    PDF_EXTRACTION_AVAILABLE = True
except ImportError:
    PDF_EXTRACTION_AVAILABLE = False
    print("⚠️ PyPDF2 not available - PDF text extraction disabled")

try:
    from docx import Document as DocxDocument
    DOCX_EXTRACTION_AVAILABLE = True
except ImportError:
    DOCX_EXTRACTION_AVAILABLE = False
    print("⚠️ python-docx not available - DOCX text extraction disabled")


class DriveIngestionService:
    """Service to sync Google Drive files to database"""

    def __init__(self, drive_client: GoogleDriveClient, db: Session):
        self.drive_client = drive_client
        self.db = db
        self.tagger = SimpleTagger()
        self._current_user_email = None
        
        # Create temp_downloads directory if it doesn't exist
        temp_dir = 'temp_downloads'
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir, exist_ok=True)
            print(f"✅ Created directory: {temp_dir}")

    def sync_all_files(self) -> Dict:
        """
        Sync all files from Google Drive to database
        ✅ Each user's files stored with their email
        ✅ ALWAYS creates tags for all files
        ✅ NEW: Auto-detects templates from filename
        """
        print("🔄 Starting full sync from Google Drive (ALWAYS CREATE TAGS)...")

        # ✅ GET CURRENT USER EMAIL ONCE AT START
        current_user_email = None
        try:
            about = self.drive_client.service.about().get(fields='user').execute()
            current_user_email = about['user']['emailAddress']
            self._current_user_email = current_user_email
            print(f"📧 Syncing for user: {current_user_email}")
        except Exception as e:
            print(f"⚠️ Could not get user: {e}")

        stats = {
            "total_files": 0,
            "new_files": 0,
            "updated_files": 0,
            "templates_detected": 0,  # NEW: Count of template files
            "tags_created": 0,
            "errors": 0,
            "skipped": 0
        }

        try:
            page_token = None

            while True:
                results = self.drive_client.list_files(page_size=100, page_token=page_token)
                files = results.get('files', [])

                if not files:
                    break

                stats["total_files"] += len(files)

                for file in files:
                    try:
                        # ✅ PASS EMAIL TO PROCESS FILE - ALWAYS CREATES TAGS
                        result = self._process_file(file, current_user_email)
                        
                        if result == "new":
                            stats["new_files"] += 1
                            stats["tags_created"] += 1
                        elif result == "updated_tags":
                            stats["updated_files"] += 1
                            stats["tags_created"] += 1
                        elif result == "skipped":
                            stats["skipped"] += 1
                        
                        # Count templates
                        if result in ["new", "updated_tags"]:
                            file_name = file.get('name', '').lower()
                            if any(keyword in file_name for keyword in ['template', 'templates']):
                                stats["templates_detected"] += 1
                                
                    except Exception as e:
                        print(f"❌ Error processing file: {e}")
                        stats["errors"] += 1

                page_token = results.get('nextPageToken')
                if not page_token:
                    break

            self._save_checkpoint(stats)
            print(f"🎉 Sync complete: {stats}")
            print(f"   📊 Tags created for {stats['tags_created']}/{stats['total_files']} files")
            print(f"   📋 Templates detected: {stats['templates_detected']}")
            return stats

        except Exception as e:
            print(f"❌ Sync failed: {e}")
            import traceback
            print(traceback.format_exc())
            return stats

    def _process_file(self, file_data: Dict, account_email: str = None) -> str:
        """
        Process a single file - ALWAYS CREATE TAGS, even for unchanged files
        ✅ NEW: Detects templates from filename and logs them
        """
        drive_file_id = file_data.get('id')
        file_name = file_data.get('name', 'Unknown')
        
        # ✅ LOG TEMPLATE DETECTION
        file_name_lower = file_name.lower()
        if any(keyword in file_name_lower for keyword in ['template', 'templates']):
            print(f"📋 Processing TEMPLATE file: {file_name}")
        else:
            print(f"📄 Processing file: {file_name}")
        
        # Get modified time from Google Drive (but we'll ignore it for tagging)
        modified_time_str = file_data.get('modifiedTime')
        drive_modified_at = None
        
        if modified_time_str:
            try:
                modified_time_str = modified_time_str.replace('Z', '+00:00')
                drive_modified_at = datetime.fromisoformat(modified_time_str)
                
                if drive_modified_at.tzinfo is None:
                    drive_modified_at = drive_modified_at.replace(tzinfo=timezone.utc)
                
                drive_modified_at = drive_modified_at.replace(microsecond=0)
                
            except Exception as e:
                print(f"⚠️ Could not parse time for {file_name}: {e}")
                drive_modified_at = datetime.now(timezone.utc)
    
        try:
            # Check if file exists in database
            existing_doc = self.db.query(Document).filter(
                Document.drive_file_id == drive_file_id
            ).first()
            
            if existing_doc:
                print(f"🔄 Updating existing document: {file_name}")
                
                # Check if file was modified (only for metadata update)
                db_modified_at = existing_doc.modified_at
                
                if db_modified_at and db_modified_at.tzinfo is None:
                    db_modified_at = db_modified_at.replace(tzinfo=timezone.utc)
                
                if db_modified_at:
                    db_modified_at = db_modified_at.replace(microsecond=0)
                
                # Update metadata if file was modified
                if drive_modified_at and db_modified_at and drive_modified_at > db_modified_at:
                    print(f"📝 File modified, updating metadata: {file_name}")
                    file_metadata = self._extract_metadata(file_data, account_email)

                    # Update other fields (everything except content_type)
                    for key, value in file_metadata.items():
                        if key != 'content_type':
                            setattr(existing_doc, key, value)

                    existing_doc.db_updated_at = datetime.utcnow()
                    self.db.commit()

                # ⭐ Re-classify content_type from the *current* filename so
                # renames in Drive (e.g. "Foo.docx" → "Foo_Template.docx")
                # are reflected on the next sync. We only ever upgrade
                # to/from TEMPLATE-vs-OTHER based on filename keywords —
                # explicit categories set by other code (CLAUSE_SET,
                # PRACTICE_NOTE, KNOWLEDGE_MATERIAL, etc.) are preserved.
                self._reclassify_template_status(existing_doc, file_name)

                # Re-tag ONLY when there's a reason to: either the file
                # has never been indexed, or it was modified in Drive
                # since the last index. Re-tagging on every sync used to
                # silently undo user edits — e.g. a tag the user removed
                # in the UI would be re-added by `_save_tags_to_database`
                # because it rebuilds `source='content_analysis'` rows
                # from scratch every run.
                needs_retag = (
                    not existing_doc.last_indexed_at
                    or (
                        drive_modified_at
                        and existing_doc.modified_at
                        and drive_modified_at > existing_doc.modified_at
                    )
                )
                if needs_retag:
                    print(f"🏷️  Creating content-based tags for: {file_name}")
                    self._create_simple_tags(drive_file_id, file_data)
                    self._queue_processing_tasks(drive_file_id)
                else:
                    print(f"⏭️  Skipping tag regeneration (unchanged): {file_name}")

                return "updated_tags" if needs_retag else "skipped"
                
            else:
                # New file
                print(f"✅ Adding new document: {file_name}")
                file_metadata = self._extract_metadata(file_data, account_email)
                
                new_doc = Document(**file_metadata)
                self.db.add(new_doc)
                self.db.flush()  # Get the ID
                self.db.commit()
                
                # ⭐⭐⭐ CREATE TAGS FOR NEW FILE ⭐⭐⭐
                print(f"🏷️  Creating content-based tags for new file: {file_name}")
                self._create_simple_tags(drive_file_id, file_data)
                
                self._queue_processing_tasks(drive_file_id)
                
                print(f"✅ Added: {file_name} (User: {account_email})")
                return "new"
        
        except Exception as e:
            print(f"❌ Error processing '{file_name}': {str(e)}")
            import traceback
            print(traceback.format_exc())
            raise

    # Categories we manage automatically from the filename. Anything
    # outside this set (e.g. CLAUSE_SET, PRACTICE_NOTE, KNOWLEDGE_MATERIAL)
    # was set deliberately by another code path and must not be clobbered
    # by a Drive sync.
    _FILENAME_MANAGED_TYPES = {ContentType.TEMPLATE, ContentType.OTHER, None}

    def _reclassify_template_status(self, doc: Document, file_name: str) -> None:
        """
        Keep a document's content_type in sync with its current Drive
        filename. Handles the two cases the original sync missed:

          * File was first ingested under a non-template name (stored as
            OTHER) and later renamed to include "template" → upgrade.
          * File was first ingested with "template" in its name and later
            renamed to remove it → downgrade to OTHER.

        Only documents that are currently TEMPLATE / OTHER / NULL are
        touched — categories like PRACTICE_NOTE that were assigned by
        Note_controller stay put.
        """
        try:
            current = doc.content_type
            if current not in self._FILENAME_MANAGED_TYPES:
                return

            looks_like_template = any(
                kw in (file_name or "").lower()
                for kw in ("template", "templates")
            )
            target = ContentType.TEMPLATE if looks_like_template else ContentType.OTHER

            if current == target:
                return

            doc.content_type = target
            self.db.commit()
            print(
                f"♻️  Reclassified {file_name}: "
                f"{current.value if current else 'NULL'} → {target.value}"
            )
        except Exception as e:
            # Reclassification is opportunistic — never let it break the
            # surrounding sync.
            self.db.rollback()
            print(f"⚠️ Reclassify failed for {file_name}: {e}")

    def _create_simple_tags(self, document_id: str, file_data: Dict):
        """
        Create tags based on document CONTENT
        Creates tags in database ONLY when found in content
        """
        # Get document from database
        doc = self.db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            doc = self.db.query(Document).filter(Document.drive_file_id == document_id).first()
        if not doc:
            print(f"⚠️ Skipping tag creation; document not in DB: {document_id}")
            return
        
        file_name = file_data.get('name', '')
        
        print(f"🔍 Analyzing content for tags: {file_name}")
        
        # STEP 1: Extract text from document
        document_text = self._extract_document_text(doc, file_name)
        
        if not document_text or len(document_text) < 10:
            print(f"ℹ️  No text content found for: {file_name}")
            # Use filename as text for tagging
            document_text = file_name
        
        # STEP 2: Generate tags from content
        tags_to_create = self.tagger.generate_tags(
            file_name=file_name,
            mime_type=doc.mime_type,
            document_text=document_text
        )
        
        if not tags_to_create:
            print(f"ℹ️  No tags found for: {file_name}")
            return
        
        print(f"🏷️  Creating content-based tags for {file_name}: {tags_to_create}")
        
        # STEP 3: Save tags to database
        self._save_tags_to_database(doc.id, tags_to_create)
    
    def _extract_document_text(self, doc, file_name: str) -> str:
        """Extract text from document file"""
        document_text = ""
        
        # Try to get document text from extracted text file
        if hasattr(doc, 'derived_text_path') and doc.derived_text_path:
            try:
                with open(doc.derived_text_path, 'r', encoding='utf-8') as f:
                    document_text = f.read()
                print(f"📖 Found extracted text: {len(document_text)} characters")
                return document_text
            except Exception as e:
                print(f"⚠️ Could not read extracted text: {e}")
        
        # If no extracted text, try to extract from temp file
        temp_file_path = os.path.join('temp_downloads', f"{doc.id}.temp")
        if os.path.exists(temp_file_path):
            try:
                # Simple text extraction
                if doc.mime_type == 'application/pdf' and PDF_EXTRACTION_AVAILABLE:
                    with open(temp_file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                document_text += page_text + "\n"
                
                elif doc.mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                      'application/msword'] and DOCX_EXTRACTION_AVAILABLE:
                    docx = DocxDocument(temp_file_path)
                    for para in docx.paragraphs:
                        if para.text:
                            document_text += para.text + "\n"
                
                elif doc.mime_type == 'text/plain':
                    with open(temp_file_path, 'r', encoding='utf-8') as f:
                        document_text = f.read()
                
                print(f"📖 Extracted text directly: {len(document_text)} characters")
                
            except Exception as e:
                print(f"⚠️ Could not extract text: {e}")
        
        return document_text
    
    def _save_tags_to_database(self, document_id: str, tags_to_create: List[str]):
        """
        Save content-analysis tags to the database.

        Behavior:
          * Tombstone rows (`source='user_removed'`) are *preserved* —
            those represent tags the user explicitly removed in the UI
            and must not be silently re-added.
          * Existing `source='content_analysis'` rows are wiped and
            rebuilt from the latest content.
          * User-applied tags (`source='user'`) are left untouched.
        """
        try:
            # 1) Build the per-document blocklist from existing tombstones.
            #    Anything in here will be skipped no matter how many times
            #    the auto-tagger thinks it should be applied.
            blocked_rows = self.db.query(DocumentTag.tag_id).filter(
                DocumentTag.document_id == document_id,
                DocumentTag.source == "user_removed",
            ).all()
            blocked_tag_ids = {row[0] for row in blocked_rows}

            # 2) Wipe ONLY the previous content-analysis links — leave
            #    user-applied and user-removed rows alone.
            existing_tags = self.db.query(DocumentTag).filter(
                DocumentTag.document_id == document_id,
                DocumentTag.source == 'content_analysis'
            ).all()
            for tag in existing_tags:
                self.db.delete(tag)

            tags_added = 0
            for tag_name in set(tags_to_create):
                if not self._is_tag_in_master_taxonomy(tag_name):
                    print(f"   ⏭️ Skipping: {tag_name} (not in master taxonomy)")
                    continue

                # Find or create tag in database
                tag = self.db.query(Tag).filter(Tag.name == tag_name).first()
                if not tag:
                    category = "custom"
                    if ": " in tag_name:
                        category = tag_name.split(": ")[0]

                    tag = Tag(
                        name=tag_name,
                        category=category,
                        created_at=datetime.utcnow(),
                        updated_at=datetime.utcnow()
                    )
                    self.db.add(tag)
                    self.db.flush()  # Get the ID
                    print(f"   📝 Created new tag: {tag_name}")

                # Honor the user's removal — never re-add a tombstoned tag.
                if tag.id in blocked_tag_ids:
                    print(f"   🚫 Skipping (user-removed): {tag_name}")
                    continue

                doc_tag = DocumentTag(
                    document_id=document_id,
                    tag_id=tag.id,
                    confidence_score=1.0,
                    source='content_analysis',
                    created_at=datetime.utcnow()
                )
                self.db.add(doc_tag)
                tags_added += 1
                print(f"   ✅ Added tag: {tag_name}")
            
            if tags_added > 0:
                self.db.commit()
                print(f"🎉 Added {tags_added} tags to database")
            else:
                print(f"ℹ️  No valid tags to save")
            
        except Exception as e:
            self.db.rollback()
            print(f"❌ Error saving tags: {e}")
            import traceback
            print(traceback.format_exc())
    
    def _is_tag_in_master_taxonomy(self, tag_name: str) -> bool:
        """Check if a tag is in the master taxonomy"""
        from tagging import ContentBasedTagger
        
        # Remove category prefix for checking
        check_name = tag_name
        if ": " in tag_name:
            check_name = tag_name.split(": ")[1]
        
        # Check all categories in master taxonomy
        for category, tag_dict in ContentBasedTagger.MASTER_TAXONOMY.items():
            if check_name in tag_dict:
                return True
        
        return False

    def _extract_metadata(self, file_data: Dict, account_email: str = None) -> Dict:
        """
        Extract metadata from Google Drive file data
        ✅ NEW: Include account_email to track which user this file belongs to
        ✅ NEW: Detect template from filename and set content_type
        """
        # Parse timestamps
        created_time = file_data.get('createdTime')
        modified_time = file_data.get('modifiedTime')

        def parse_iso_time(time_str):
            if not time_str:
                return None
            try:
                time_str = time_str.replace('Z', '+00:00')
                return datetime.fromisoformat(time_str)
            except Exception as e:
                print(f"⚠️ Could not parse time '{time_str}': {e}")
                return None

        created_at = parse_iso_time(created_time)
        modified_at = parse_iso_time(modified_time)

        # Extract owner info
        owners = file_data.get('owners', [])
        owner_email = owners[0].get('emailAddress') if owners else None
        owner_name = owners[0].get('displayName') if owners else None

        # Get file extension
        file_name = file_data.get('name', '')
        file_format = self._get_file_extension(file_name)

        # ✅ DETECT CONTENT TYPE FROM FILENAME
        file_name_lower = file_name.lower()
        
        # Check for template keywords in filename
        if any(keyword in file_name_lower for keyword in ['template', 'templates']):
            content_type = ContentType.TEMPLATE
            print(f"   📋 Detected as TEMPLATE: {file_name}")
        else:
            content_type = ContentType.OTHER
            print(f"   📄 Detected as OTHER: {file_name}")

        # Create checksum
        checksum = hashlib.md5(
            f"{file_data.get('id')}{modified_time}".encode()
        ).hexdigest()

        metadata = {
            'id': file_data.get('id'),
            'drive_file_id': file_data.get('id'),
            'title': file_name,
            'mime_type': file_data.get('mimeType'),
            'file_format': file_format,
            'size_bytes': int(file_data.get('size', 0)) if file_data.get('size') else None,
            'owner_email': owner_email,
            'owner_name': owner_name,
            'file_url': file_data.get('webViewLink'),
            'thumbnail_link': file_data.get('thumbnailLink'),
            'icon_link': file_data.get('iconLink'),
            # ✅ ADD CONTENT TYPE TO METADATA
            'content_type': content_type,
            'created_at': created_at,
            'modified_at': modified_at,
            'status': 'active',
            # ✅ THIS IS THE KEY - Store which account this file belongs to
            'account_email': account_email,
            'account_id': None,
            'checksum': checksum,
            'db_created_at': datetime.utcnow(),
            'db_updated_at': datetime.utcnow()
        }

        return metadata

    def _get_file_extension(self, filename: str) -> Optional[str]:
        """Extract file extension"""
        if '.' in filename:
            return '.' + filename.rsplit('.', 1)[1].lower()
        return None

    def _queue_processing_tasks(self, document_id: str):
        """
        Queue processing tasks for a document
        Creates 3 tasks per document:
        1. Extract text
        2. AI tagging
        3. Create embeddings
        """
        tasks = [
            TaskType.EXTRACT_TEXT,
            TaskType.AI_TAGGING,
            TaskType.CREATE_EMBEDDING
        ]

        for task_type in tasks:
            # Check if task already exists
            existing_task = self.db.query(ProcessingQueue).filter(
                ProcessingQueue.document_id == document_id,
                ProcessingQueue.task_type == task_type,
                ProcessingQueue.status.in_([ProcessingStatus.PENDING, ProcessingStatus.PROCESSING])
            ).first()

            if not existing_task:
                new_task = ProcessingQueue(
                    document_id=document_id,
                    task_type=task_type,
                    status=ProcessingStatus.PENDING,
                    priority=5 if task_type == TaskType.EXTRACT_TEXT else 7,
                    retry_count=0,
                    max_retries=3
                )
                self.db.add(new_task)

        self.db.commit()

    def _save_checkpoint(self, stats: Dict):
        """
        Save sync checkpoint
        Records when sync happened, how many files, any errors
        """
        checkpoint = SyncCheckpoint(
            source='google_drive',
            last_sync_time=datetime.utcnow(),
            files_processed=stats.get('total_files', 0),
            files_failed=stats.get('errors', 0),
            status='completed' if stats['errors'] == 0 else 'completed_with_errors'
        )
        self.db.add(checkpoint)
        self.db.commit()

    def get_last_sync_info(self) -> Optional[Dict]:
        """Get information about last sync"""
        last_checkpoint = self.db.query(SyncCheckpoint).filter(
            SyncCheckpoint.source == 'google_drive'
        ).order_by(SyncCheckpoint.created_at.desc()).first()

        if last_checkpoint:
            return {
                'last_sync_time': last_checkpoint.last_sync_time.isoformat(),
                'files_processed': last_checkpoint.files_processed,
                'files_failed': last_checkpoint.files_failed,
                'status': last_checkpoint.status
            }
        return None

    def get_sync_stats(self) -> Dict:
        """Get overall sync statistics"""
        total_docs = self.db.query(Document).count()
        
        # Count templates
        templates = self.db.query(Document).filter(
            Document.content_type == ContentType.TEMPLATE
        ).count()
        
        # Count other documents
        others = self.db.query(Document).filter(
            Document.content_type == ContentType.OTHER
        ).count()
        
        pending_tasks = self.db.query(ProcessingQueue).filter(
            ProcessingQueue.status == ProcessingStatus.PENDING
        ).count()

        processing_tasks = self.db.query(ProcessingQueue).filter(
            ProcessingQueue.status == ProcessingStatus.PROCESSING
        ).count()

        failed_tasks = self.db.query(ProcessingQueue).filter(
            ProcessingQueue.status == ProcessingStatus.FAILED
        ).count()

        return {
            'total_documents': total_docs,
            'templates': templates,
            'other_documents': others,
            'pending_tasks': pending_tasks,
            'processing_tasks': processing_tasks,
            'failed_tasks': failed_tasks,
            'last_sync': self.get_last_sync_info()
        }