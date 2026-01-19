"""
Universal Content Extractor
Extracts text content from various file types (PDF, DOCX, Google Docs, etc.)
"""
import io
import re
from typing import Optional


class UniversalContentExtractor:
    """
    Extract text content from various document types
    """
    
    def __init__(self, drive_client):
        self.drive_client = drive_client

    def extract_structured(self, file_id: str, mime_type: str):
        """
        Extract structured content (headings + paragraphs)
        """
        print(f"📄 Extracting STRUCTURED content: {file_id} ({mime_type})")

        # Google Docs
        if 'google-apps.document' in mime_type:
            return self._extract_google_doc_structured(file_id)

        # DOCX
        elif 'officedocument.wordprocessingml' in mime_type or 'msword' in mime_type:
            return self._extract_docx_structured(file_id)

        # Fallback to existing flat text logic
        text = self.extract_content(file_id, mime_type)
        return self._flat_text_to_blocks(text)

    def extract_content(self, file_id: str, mime_type: str) -> str:
        """
        Extract text content from a file
        """
        try:
            print(f"📄 Extracting content from file: {file_id} ({mime_type})")
            
            # Google Docs
            if 'google-apps.document' in mime_type:
                return self._extract_from_google_doc(file_id)
            
            # PDF
            elif 'pdf' in mime_type:
                return self._extract_from_pdf(file_id)
            
            # Word Document
            elif 'officedocument.wordprocessingml' in mime_type or 'msword' in mime_type:
                return self._extract_from_docx(file_id)
            
            # Plain text
            elif 'text' in mime_type:
                return self._extract_from_text(file_id)
            
            else:
                print(f"⚠️ Unsupported file type: {mime_type}")
                return ""
                
        except Exception as e:
            print(f"❌ Error extracting content: {e}")
            import traceback
            print(traceback.format_exc())
            return ""
    
    def _extract_from_google_doc(self, file_id: str) -> str:
        """
        Extract text from Google Doc by exporting as plain text
        """
        try:
            # Export Google Doc as plain text
            request = self.drive_client.service.files().export_media(
                fileId=file_id,
                mimeType='text/plain'
            )
            
            content = request.execute()
            
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            print(f"✅ Extracted {len(content)} characters from Google Doc")
            return content
            
        except Exception as e:
            print(f"❌ Error extracting Google Doc: {e}")
            return ""
    
    def _extract_from_pdf(self, file_id: str) -> str:
        """
        Extract text from PDF using PyPDF2
        """
        try:
            import PyPDF2
            
            # Download PDF
            request = self.drive_client.service.files().get_media(fileId=file_id)
            file_content = io.BytesIO(request.execute())
            
            # Extract text
            pdf_reader = PyPDF2.PdfReader(file_content)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            print(f"✅ Extracted {len(text)} characters from PDF")
            return text
            
        except ImportError:
            print("❌ PyPDF2 not installed. Run: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"❌ Error extracting PDF: {e}")
            return ""
    
    def _extract_from_docx(self, file_id: str) -> str:
        """
        Extract text from Word document using python-docx
        """
        try:
            import docx
            
            # Download DOCX
            request = self.drive_client.service.files().get_media(fileId=file_id)
            file_content = io.BytesIO(request.execute())
            
            # Extract text
            doc = docx.Document(file_content)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            print(f"✅ Extracted {len(text)} characters from DOCX")
            return text
            
        except ImportError:
            print("❌ python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            print(f"❌ Error extracting DOCX: {e}")
            return ""
    
    def _extract_from_text(self, file_id: str) -> str:
        """
        Extract plain text file
        """
        try:
            request = self.drive_client.service.files().get_media(fileId=file_id)
            content = request.execute()
            
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            print(f"✅ Extracted {len(content)} characters from text file")
            return content
            
        except Exception as e:
            print(f"❌ Error extracting text: {e}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text (remove extra whitespace, etc.)
        """
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()


    def _extract_google_doc_structured(self, file_id: str):
        """
        Extract headings and paragraphs from Google Docs
        """
        doc = self.drive_client.docs_service.documents().get(
            documentId=file_id
        ).execute()

        blocks = []

        for el in doc.get("body", {}).get("content", []):
            if "paragraph" not in el:
                continue

            para = el["paragraph"]
            style = para.get("paragraphStyle", {}).get("namedStyleType", "")

            text = "".join(
                r["textRun"]["content"]
                for r in para.get("elements", [])
                if "textRun" in r
            ).strip()

            if not text:
                continue

            if style.startswith("HEADING_"):
                blocks.append({
                    "type": "heading",
                    "text": text,
                    "level": int(style[-1])
                })
            else:
                blocks.append({
                    "type": "paragraph",
                    "text": text,
                    "level": None
                })

        print(f"✅ Extracted {len(blocks)} structured blocks from Google Doc")
        return blocks


    def _extract_docx_structured(self, file_id: str):
        """
        Extract headings and paragraphs from DOCX
        """
        import docx

        request = self.drive_client.service.files().get_media(fileId=file_id)
        file_content = io.BytesIO(request.execute())

        doc = docx.Document(file_content)
        blocks = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                level = None
                last_part = style.split()[-1]
                if last_part.isdigit():
                    level = int(last_part)
                blocks.append({
                    "type": "heading",
                    "text": text,
                    "level": level
                })
            else:
                blocks.append({
                    "type": "paragraph",
                    "text": text,
                    "level": None
                })

        print(f"✅ Extracted {len(blocks)} structured blocks from DOCX")
        return blocks

    def _flat_text_to_blocks(self, text: str):
        """
        Convert flat text into paragraph blocks
        """
        blocks = []

        for line in text.split("\n"):
            line = line.strip()
            if line:
                blocks.append({
                    "type": "paragraph",
                    "text": line,
                    "level": None
                })

        return blocks
